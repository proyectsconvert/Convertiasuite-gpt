"""
DOCX document processor using python-docx.
Extracts text, tables, and structure from DOCX files.
"""

import io
import logging

try:
    from docx import Document as DocxDocument
    from docx.table import Table as DocxTable
except ImportError:
    DocxDocument = None

from app.domain.entities.document import (
    DocumentType,
    ParsedContent,
    Section,
    Table,
)
from app.domain.interfaces.document_processor import IDocumentProcessor

logger = logging.getLogger(__name__)


class DocxProcessor(IDocumentProcessor):
    """
    Processes DOCX files using python-docx.
    Extracts text, tables, and document structure.
    """

    def __init__(self):
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

    @property
    def supported_type(self) -> DocumentType:
        return DocumentType.DOCX

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    async def parse(
        self,
        file_content: bytes,
        filename: str,
    ) -> ParsedContent:
        """
        Parse DOCX and extract content.
        Preserves document structure through sections.
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            full_text = ""
            sections: list[Section] = []
            tables: list[Table] = []
            current_section_title = "Introduction"
            current_section_content = ""

            # Process document content
            for para_idx, para in enumerate(doc.paragraphs):
                if not para.text.strip():
                    continue

                # Detect heading-like paragraphs
                is_heading = (
                    para.style.name.startswith("Heading") or
                    len(para.text) < 100 and para_idx > 0 and
                    all(c.isupper() or not c.isalpha() for c in para.text.split()[0])
                )

                if is_heading and current_section_content.strip():
                    # Save previous section
                    section = Section(
                        title=current_section_title,
                        content=current_section_content.strip(),
                        level=1,
                        metadata={"type": "paragraph"}
                    )
                    sections.append(section)
                    full_text += f"## {current_section_title}\n{current_section_content}\n\n"
                    current_section_content = ""

                # Update section title if heading
                if is_heading:
                    current_section_title = para.text.strip()
                else:
                    current_section_content += para.text + "\n"

            # Save final section
            if current_section_content.strip():
                section = Section(
                    title=current_section_title,
                    content=current_section_content.strip(),
                    level=1,
                    metadata={"type": "paragraph"}
                )
                sections.append(section)
                full_text += f"## {current_section_title}\n{current_section_content}\n\n"

            # Process tables
            for table_idx, docx_table in enumerate(doc.tables):
                headers = []
                rows = []

                for row_idx, row in enumerate(docx_table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]

                    if row_idx == 0:
                        headers = row_data
                    else:
                        rows.append(row_data)

                if headers:
                    table = Table(
                        headers=headers,
                        rows=rows,
                        name=f"Table_{table_idx}",
                        metadata={"index": table_idx, "type": "table"}
                    )
                    tables.append(table)
                    full_text += f"\nTable {table_idx}: {' | '.join(headers)}\n"
                    for row in rows:
                        full_text += f"{' | '.join(row)}\n"

            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(sections),
                "has_core_properties": doc.core_properties is not None,
            }

            # Add core properties if available
            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                })

            return ParsedContent(
                text=full_text.strip(),
                sections=sections,
                tables=tables,
                images=[],  # python-docx doesn't easily extract images
                metadata=metadata,
            )

        except Exception as e:
            logger.error(f"DOCX parsing error for {filename}: {str(e)}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    def get_metadata(self, file_content: bytes) -> dict:
        """Extract DOCX metadata without full parsing."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            metadata = {
                "is_valid": True,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }

            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                })

            return metadata

        except Exception as e:
            logger.error(f"DOCX metadata extraction error: {str(e)}")
            return {"is_valid": False, "error": str(e)}
