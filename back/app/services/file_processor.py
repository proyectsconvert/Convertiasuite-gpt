from io import BytesIO, StringIO
import csv
import logging
import openpyxl
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class FileProcessorService:
    @staticmethod
    def extract_text_from_pdf(contents: bytes) -> str:
        """
        Extrae el contenido de un archivo PDF página por página con marcadores estructurales.
        """
        try:
            reader = PdfReader(BytesIO(contents))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(f"--- Página {i + 1} ---\n{text.strip()}")
            
            if not pages_text:
                return "El PDF parece estar vacío o no se pudo extraer texto. Verifique si contiene solo imágenes."
            
            return "\n\n".join(pages_text)
        except Exception as e:
            logger.error(f"Error parseando PDF: {str(e)}")
            raise ValueError(f"Error al procesar el archivo PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(contents: bytes) -> str:
        """
        Extrae texto y tablas de un archivo Word (.docx), preservando el formato de tablas en Markdown.
        """
        try:
            doc = Document(BytesIO(contents))
            elements = []

            # Recorrer todos los elementos del cuerpo del documento para mantener el orden
            for element in doc.element.body:
                # Si es un párrafo
                if element.tag.endswith('p'):
                    # Buscar el párrafo correspondiente en la lista de párrafos de python-docx
                    for p in doc.paragraphs:
                        if p._element is element:
                            if p.text and p.text.strip():
                                elements.append(p.text.strip())
                            break
                # Si es una tabla
                elif element.tag.endswith('tbl'):
                    for table in doc.tables:
                        if table._element is element:
                            table_md = []
                            for row_idx, row in enumerate(table.rows):
                                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                                # Unir celdas con |
                                table_md.append("| " + " | ".join(row_cells) + " |")
                                if row_idx == 0:
                                    # Generar línea separadora para Markdown
                                    table_md.append("| " + " | ".join(["---"] * len(row_cells)) + " |")
                            if table_md:
                                elements.append("\n".join(table_md))
                            break

            # Respaldo por si el parseo ordenado falla o se saltó elementos
            if not elements:
                for p in doc.paragraphs:
                    if p.text and p.text.strip():
                        elements.append(p.text.strip())
                for table in doc.tables:
                    table_md = []
                    for row_idx, row in enumerate(table.rows):
                        row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        table_md.append("| " + " | ".join(row_cells) + " |")
                        if row_idx == 0:
                            table_md.append("| " + " | ".join(["---"] * len(row_cells)) + " |")
                    if table_md:
                        elements.append("\n".join(table_md))

            if not elements:
                return "El archivo Word (.docx) está vacío."

            return "\n\n".join(elements)
        except Exception as e:
            logger.error(f"Error parseando archivo DOCX: {str(e)}")
            raise ValueError(f"Error al procesar el archivo Word (.docx): {str(e)}")

    @staticmethod
    def analyze_tabular_data(rows: list[list]) -> str:
        """
        Analiza un conjunto de datos tabular (CSV o Excel).
        Genera un resumen estadístico detallado con métricas pre-calculadas y una vista previa
        para optimizar el contexto del LLM.
        """
        if not rows:
            return "El archivo está vacío o no contiene filas válidas."

        # Filtrar celdas vacías de la cabecera
        raw_headers = rows[0]
        headers = []
        for i, h in enumerate(raw_headers):
            if h is not None and str(h).strip() != "":
                headers.append(str(h).strip())
            else:
                headers.append(f"Columna_{i+1}")

        data_rows = rows[1:]
        total_rows = len(data_rows)
        num_cols = len(headers)

        summary = []
        summary.append("### Resumen General del Dataset")
        summary.append(f"- **Total de filas de datos**: {total_rows}")
        summary.append(f"- **Total de columnas**: {num_cols}")
        summary.append(f"- **Nombres de columnas**: {', '.join([f'`{h}`' for h in headers])}")
        summary.append("")

        summary.append("### Análisis de Columnas")

        for col_idx in range(num_cols):
            col_name = headers[col_idx]
            col_values = []
            for r in data_rows:
                if col_idx < len(r):
                    val = r[col_idx]
                    if val is not None and str(val).strip() != "":
                        col_values.append(val)

            non_null_count = len(col_values)
            null_count = total_rows - non_null_count

            if non_null_count == 0:
                summary.append(f"#### Columna: `{col_name}`")
                summary.append(f"- **Estado**: Todos los valores son nulos/vacíos (100% nulos)")
                summary.append("")
                continue

            # Determinar si la columna es mayoritariamente numérica
            numeric_values = []
            for v in col_values:
                try:
                    clean_v = str(v).replace("$", "").replace(",", "").replace("%", "").strip()
                    numeric_values.append(float(clean_v))
                except ValueError:
                    pass

            is_numeric = len(numeric_values) >= 0.8 * non_null_count

            summary.append(f"#### Columna: `{col_name}`")
            summary.append(f"- **Valores no nulos (llenos)**: {non_null_count} / {total_rows} ({non_null_count/total_rows*100:.1f}%)")
            summary.append(f"- **Valores nulos/vacíos**: {null_count} ({null_count/total_rows*100:.1f}%)")

            if is_numeric:
                col_min = min(numeric_values)
                col_max = max(numeric_values)
                col_sum = sum(numeric_values)
                col_mean = col_sum / len(numeric_values)

                def format_num(val):
                    if val.is_integer():
                        return f"{int(val)}"
                    return f"{val:.2f}"

                summary.append(f"- **Tipo**: Numérico")
                summary.append(f"- **Estadísticas Calculadas**:")
                summary.append(f"  - Mínimo: {format_num(col_min)}")
                summary.append(f"  - Máximo: {format_num(col_max)}")
                summary.append(f"  - Suma Total: {format_num(col_sum)}")
                summary.append(f"  - Promedio: {format_num(col_mean)}")
            else:
                summary.append(f"- **Tipo**: Categórico / Texto")
                frequencies = {}
                for v in col_values:
                    str_v = str(v).strip()
                    frequencies[str_v] = frequencies.get(str_v, 0) + 1

                unique_count = len(frequencies)
                summary.append(f"- **Valores únicos**: {unique_count}")

                sorted_freqs = sorted(frequencies.items(), key=lambda x: x[1], reverse=True)
                top_n = sorted_freqs[:10]
                summary.append("- **Distribución de frecuencias (Top 10)**:")
                for val, count in top_n:
                    summary.append(f"  - `{val}`: {count} ocurrencias ({count/total_rows*100:.1f}%)")

            summary.append("")

        # Vista previa de las primeras 5 filas
        summary.append("### Vista Previa de los Datos (Primeras 5 filas)")
        table_header = "| " + " | ".join(headers) + " |"
        table_separator = "| " + " | ".join(["---"] * num_cols) + " |"
        summary.append(table_header)
        summary.append(table_separator)

        for r in data_rows[:5]:
            row_cells = []
            for col_idx in range(num_cols):
                val = r[col_idx] if col_idx < len(r) else ""
                row_cells.append(str(val).strip() if val is not None else "")
            summary.append("| " + " | ".join(row_cells) + " |")

        return "\n".join(summary)

    @classmethod
    def extract_text_from_csv(cls, contents: bytes) -> str:
        """
        Extrae datos tabulares de un archivo CSV y genera el análisis estadístico.
        """
        try:
            text_data = contents.decode("utf-8", errors="replace")
            reader = csv.reader(StringIO(text_data))
            rows = []
            for row in reader:
                if any(cell is not None and str(cell).strip() != "" for cell in row):
                    rows.append(row)
            return cls.analyze_tabular_data(rows)
        except Exception as e:
            logger.error(f"Error parseando CSV: {str(e)}")
            raise ValueError(f"Error al procesar el archivo CSV: {str(e)}")

    @classmethod
    def extract_text_from_excel(cls, contents: bytes) -> str:
        """
        Extrae datos tabulares de un archivo Excel (.xlsx) y genera el análisis estadístico.
        """
        try:
            wb = openpyxl.load_workbook(BytesIO(contents), data_only=True, read_only=True)
            sheet = wb.active
            rows = []
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None and str(cell).strip() != "" for cell in row):
                    rows.append(list(row))
            return cls.analyze_tabular_data(rows)
        except Exception as e:
            logger.error(f"Error parseando Excel: {str(e)}")
            raise ValueError(f"Error al procesar el archivo Excel: {str(e)}")
