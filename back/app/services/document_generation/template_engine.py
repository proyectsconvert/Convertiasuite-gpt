from __future__ import annotations
import os
import re
import logging
from typing import TYPE_CHECKING, Any, Dict, List
import pathlib
from jinja2 import Environment, FileSystemLoader, select_autoescape
from app.domain.entities.document_content import DocumentContent, Section, TableData
from app.core.files_config import BRAND_CONFIG, BASE_DIR
if TYPE_CHECKING:
    from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

# Rutas de recursos
_TEMPLATES_DIR = os.path.join(BASE_DIR, "assets", "templates")
_LOGOS_DIR = os.path.join(BASE_DIR, "assets", "logos")

class TemplateEngine:

    def __init__(self, brand: str = "convertia"):
        self._brand = brand
        self._brand_cfg = BRAND_CONFIG.get(brand, BRAND_CONFIG["convertia"])
        self._jinja_env = Environment(
            loader=FileSystemLoader(_TEMPLATES_DIR),
            autoescape=select_autoescape(["html"]),
        )

    def build_docx_context(
        self,
        content: DocumentContent,
        doc: "DocxTemplate",
    ) -> Dict[str, Any]:
        try:
            from docxtpl import InlineImage
            from docx.shared import Inches

            logo_path = self._brand_cfg["logos"].get("docs", "")
            logo = (
                InlineImage(doc, logo_path, width=Inches(1.5))
                if logo_path and os.path.exists(logo_path)
                else None
            )
        except Exception as e:
            logger.warning(f"No se pudo cargar el logo para DOCX: {e}")
            logo = None

        return {
            # Metadata del documento
            "title": content.title,
            "subtitle": content.subtitle or "",
            "author": content.author,
            "date": content.get_date(),
            "classification": content.classification,
            "logo": logo,

            # Secciones jerárquicas
            "sections": [self._section_to_dict(s) for s in content.sections],
            "sections_h1": [self._section_to_dict(s) for s in content.sections_by_level(1)],
            "sections_h2": [self._section_to_dict(s) for s in content.sections_by_level(2)],
            "sections_h3": [self._section_to_dict(s) for s in content.sections_by_level(3)],

            # Tablas globales
            "tables": [self._table_to_dict(t) for t in content.tables],

            # Branding
            "brand_name": self._brand_cfg.get("nametag", "Convertia"),
            "brand_colors": self._brand_cfg.get("colors", {}),
        }

    # PDF — contexto HTML para WeasyPrint

    def build_html_context(self, content: DocumentContent) -> str:
        try:
            template = self._jinja_env.get_template("report_template.html", "word_template.html")
        except Exception as e:
            logger.warning(f"Plantilla HTML no encontrada, usando template inline: {e}")
            return self._render_inline_html(content)

        logo_main_path = self._brand_cfg["logos"].get("main", "")
        logo_main_src = pathlib.Path(logo_main_path).as_uri() if logo_main_path and os.path.exists(logo_main_path) else ""

        logo_docs_path = self._brand_cfg["logos"].get("docs", "")
        logo_docs_src = pathlib.Path(logo_docs_path).as_uri() if logo_docs_path and os.path.exists(logo_docs_path) else ""

        words = content.title.split()
        if len(words) > 1:
            title_part1 = " ".join(words[:-1]) + " "
            title_part2 = words[-1]
        else:
            title_part1 = content.title
            title_part2 = ""

        context = {
            "title": content.title,
            "title_part1": title_part1,
            "title_part2": title_part2,
            "subtitle": content.subtitle or "",
            "author": content.author,
            "date": content.get_date(),
            "classification": content.classification,
            "sections": [self._section_to_dict(s) for s in content.sections],
            "tables": [self._table_to_dict(t) for t in content.tables],
            "logo_src": logo_docs_src,
            "logo_main_src": logo_main_src,
            "logo_docs_src": logo_docs_src,
            "brand_name": self._brand_cfg.get("nametag", "Convertia"),
            "colors": self._brand_cfg.get("colors", {}),
        }

        return template.render(**context)
    #ppt
    def build_pptx_context(self, content: DocumentContent) -> Dict[str, Any]:
      
        slides = []

        # Slide de portada
        slides.append({
            "type": "cover",
            "title": content.title,
            "subtitle": content.subtitle or content.author,
            "date": content.get_date(),
            "classification": content.classification,
        })

        # Slide de índice / contenido
        index_items = []
        h1_count = 0
        for section in content.sections:
            if section.level == 1:
                h1_count += 1
                index_items.append({
                    "number": f"{h1_count:02d}",
                    "title": section.title,
                    "subsections": []
                })
            elif section.level == 2 and index_items:
                sub_count = len(index_items[-1]["subsections"]) + 1
                index_items[-1]["subsections"].append({
                    "number": f"{sub_count}",
                    "title": section.title
                })

        slides.append({
            "type": "index",
            "items": index_items
        })

        # Slides de contenido por sección
        section_count = 0
        for section in content.sections:
            if section.level == 1:
                section_count += 1
                slides.append({
                    "type": "section",
                    "number": section_count,
                    "title": section.title,
                    "subtitle": section.content[:120] if section.content else "",
                })
            elif section.level == 2:
                slide_data: Dict[str, Any] = {
                    "type": "content",
                    "title": section.title,
                    "content": section.content,
                    "bullets": section.bullets or [],
                }
                if section.table:
                    slide_data["table"] = self._table_to_dict(section.table)
                slides.append(slide_data)
            elif section.level == 3:
                # H3 se agrega al slide anterior si existe
                if slides and slides[-1]["type"] == "content":
                    prev = slides[-1]
                    prev["subheadings"] = prev.get("subheadings", [])
                    prev["subheadings"].append({
                        "title": section.title,
                        "content": section.content,
                        "bullets": section.bullets or [],
                    })
                else:
                    slides.append({
                        "type": "content",
                        "title": section.title,
                        "content": section.content,
                        "bullets": section.bullets or [],
                    })

        # Tablas globales como slides dedicados
        for table in content.tables:
            slides.append({
                "type": "table",
                "title": table.caption or "Datos",
                "table": self._table_to_dict(table),
            })

        # Slide de cierre
        slides.append({"type": "closing"})

        return {
            "slides": slides,
            "brand_name": self._brand_cfg.get("nametag", "Convertia"),
            "logo_path": self._brand_cfg["logos"].get("white", ""),
        }

  
    def build_excel_context(self, content: DocumentContent) -> Dict[str, Any]:
        sheets = []

        for i, table in enumerate(content.tables):
            sheets.append({
                "name": table.caption or f"Datos {i + 1}",
                "headers": table.headers,
                "rows": table.rows,
                "summary": table.summary,
            })

        for section in content.sections:
            if section.table:
                sheets.append({
                    "name": section.title[:31],  
                    "headers": section.table.headers,
                    "rows": section.table.rows,
                    "summary": section.table.summary,
                })

        if not sheets:
            sheets.append({
                "name": "Resumen",
                "headers": ["Sección", "Contenido"],
                "rows": [[s.title, s.content[:500]] for s in content.sections],
                "summary": None,
            })

        return {
            "title": content.title,
            "sheets": sheets,
            "excel_cfg": self._brand_cfg.get("excel", {}),
        }

    @staticmethod
    def _section_to_dict(section: Section) -> Dict[str, Any]:
        paragraphs = []
        if section.content:
            for para in section.content.split("\n"):
                if para.strip():
                    p_text = para.strip()
                    if p_text.startswith(">"):
                        # Callout box! Parse simple **bold** tags to class span
                        callout_text = p_text.lstrip(">").strip()
                        parts = re.split(r"(\*\*.*?\*\*)", callout_text)
                        html_parts = []
                        for part in parts:
                            if part.startswith("**") and part.endswith("**"):
                                html_parts.append(
                                    f'<span class="callout-highlight">{part[2:-2]}</span>'
                                )
                            else:
                                html_parts.append(part)
                        callout_html = "".join(html_parts)
                        paragraphs.append({"type": "callout", "text": callout_html})
                    else:
                        paragraphs.append({"type": "text", "text": p_text})
        return {
            "title": section.title,
            "level": section.level,
            "content": section.content,
            "paragraphs": paragraphs,
            "bullets": section.bullets or [],
            "has_bullets": bool(section.bullets),
            "table": (
                TemplateEngine._table_to_dict(section.table)
                if section.table else None
            ),
            "has_table": section.table is not None,
        }

    @staticmethod
    def _table_to_dict(table: TableData) -> Dict[str, Any]:
        rows_processed = []
        for row in table.rows:
            is_highlight = any(str(val).strip().lower() == "destacado" for val in row)
            rows_processed.append({
                "cells": [str(val) for val in row],
                "is_highlight": is_highlight
            })
        return {
            "headers": table.headers,
            "rows": rows_processed,
            "caption": table.caption or "",
            "row_count": table.row_count,
            "column_count": table.column_count,
        }


    def _render_inline_html(self, content: DocumentContent) -> str:
        colors = self._brand_cfg.get("colors", {})
        primary = colors.get("primary", "#011E23")
        accent = colors.get("accent_green", "#1AEB9F")
        green = colors.get("secondary_dark_green", "#10473f")
        muted = "#718096"

        words = content.title.split()
        if len(words) > 1:
            title_part1 = " ".join(words[:-1]) + " "
            title_part2 = words[-1]
        else:
            title_part1 = content.title
            title_part2 = ""

        sections_html = ""
        for section in content.sections:
            if section.level == 1:
                sections_html += f'<h1 class="section-divider">{section.title}</h1>\n'
            elif section.level == 2:
                sections_html += f'<h2 style="border-left: 3px solid {accent}; padding-left: 10px; color: {green};">{section.title}</h2>\n'
            else:
                sections_html += f'<h3 style="color: #2D3748;">{section.title}</h3>\n'

            if section.content:
                for para in section.content.split("\n"):
                    if para.strip():
                        p_text = para.strip()
                        if p_text.startswith(">"):
                            callout_text = p_text.lstrip(">").strip()
                            parts = re.split(r"(\*\*.*?\*\*)", callout_text)
                            html_parts = []
                            for part in parts:
                                if part.startswith("**") and part.endswith("**"):
                                    html_parts.append(
                                        f'<span style="color: {accent};">{part[2:-2]}</span>'
                                    )
                                else:
                                    html_parts.append(part)
                            callout_html = "".join(html_parts)
                            sections_html += (
                                f'<div class="callout-box"><p style="color: white; text-align: center; '
                                f'font-weight: bold; margin: 0;">{callout_html}</p></div>\n'
                            )
                        else:
                            sections_html += f"<p>{p_text}</p>\n"
            if section.bullets:
                sections_html += "<ul>\n"
                for bullet in section.bullets:
                    sections_html += f"  <li>{bullet}</li>\n"
                sections_html += "</ul>\n"
            if section.table:
                sections_html += self._table_to_html(section.table)

        tables_html = "".join(self._table_to_html(t) for t in content.tables)

        logo_path = self._brand_cfg["logos"].get("main", "")
        logo_html = ""
        if logo_path and os.path.exists(logo_path):
            valid_uri = pathlib.Path(logo_path).as_uri()
            logo_html = f'<div style="text-align: right; margin-bottom: 80pt;"><img src="{valid_uri}" style="height: 28px; object-fit: contain;" /></div>'

        return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>{content.title}</title>
<style>
  @page {{
    size: letter;
    margin: 2.4cm 1.8cm 2.2cm 1.8cm;
    @bottom-left {{ content: "© Intelligence Customer Acquisition — Convertia — Documentación interna"; font-family: Inter, Arial, sans-serif; font-size: 7.5pt; color: #718096; }}
    @bottom-right {{ content: "Pág. " counter(page); font-family: Inter, Arial, sans-serif; font-size: 7.5pt; color: #718096; }}
  }}
  body {{ font-family: Inter, Arial, sans-serif; font-size: 10.5pt; color: #2d3748; line-height: 1.65; }}
  .cover {{ page-break-after: always; padding-top: 80pt; position: relative; }}
  .cover h1 {{ font-size: 32pt; color: {primary}; margin-bottom: 10pt; font-weight: 700; line-height: 1.15; }}
  .cover h1 span.title-highlight {{ color: {muted}; }}
  .cover-badge {{ display: inline-block; background-color: #f5f5f5; color: {muted}; font-size: 8pt; font-weight: bold; padding: 4px 10px; text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 12pt; }}
  .cover .subtitle {{ font-size: 13pt; color: {green}; margin-bottom: 60pt; }}
  .cover .meta {{ margin-top: 60pt; font-size: 9pt; color: #718096; }}
  .section-divider {{ page-break-before: always; background: #ffffff; color: {primary}; font-size: 26pt; font-weight: 700; padding: 28pt 0 22pt; margin: 0 0 24pt; border-bottom: 3pt solid {accent}; }}
  p {{ margin-bottom: 7pt; text-align: justify; color: {primary}; }}
  ul {{ padding-left: 18pt; margin: 6pt 0 10pt; }}
  li {{ margin-bottom: 4pt; }}
  ul li::marker {{ color: {accent}; }}
  .callout-box {{ background-color: {primary}; color: #ffffff; font-weight: bold; padding: 16px 20px; margin: 14pt 0 18pt; text-align: center; }}
  table {{ width: 100%; border-collapse: collapse; margin: 14pt 0 18pt; font-size: 9pt; }}
  thead th {{ background: {primary}; color: white; padding: 7pt 9pt; text-align: left; font-weight: 600; }}
  tbody tr:nth-child(even) {{ background: #f5f5f5; }}
  tbody tr:nth-child(odd) {{ background: #ffffff; }}
  tbody tr.row-highlight {{ background-color: #E6FFFA !important; color: {green}; font-weight: bold; }}
  tbody tr.row-highlight td {{ color: {green}; font-weight: bold; }}
  tbody td {{ padding: 5pt 9pt; border-bottom: 0.5pt solid #dbdffc; vertical-align: top; }}
</style>
</head>
<body>
  <div class="cover">
    {logo_html}
    <h1>{title_part1}{f'<span class="title-highlight">{title_part2}</span>' if title_part2 else ''}</h1>
    <div class="subtitle">{content.subtitle or content.author}</div>
    <div class="meta">
      <strong>Fecha:</strong> {content.get_date()}<br>
      <strong>Clasificación:</strong> {content.classification}
    </div>
  </div>
  {sections_html}
  {tables_html}
</body>
</html>"""

    @staticmethod
    def _table_to_html(table: TableData) -> str:
        header_cells = "".join(f"<th>{h}</th>" for h in table.headers)
        rows_html = ""
        for row in table.rows:
            is_highlight = any(str(val).strip().lower() == "destacado" for val in row)
            tr_class = ' class="row-highlight"' if is_highlight else ""
            cells = "".join(f"<td>{cell}</td>" for cell in row)
            rows_html += f"<tr{tr_class}>{cells}</tr>\n"
        caption = f"<caption>{table.caption}</caption>" if table.caption else ""
        return f"""<table>{caption}
  <thead><tr>{header_cells}</tr></thead>
  <tbody>{rows_html}</tbody>
</table>"""
