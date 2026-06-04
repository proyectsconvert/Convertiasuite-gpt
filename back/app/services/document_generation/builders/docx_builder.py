import os
import logging
import re
from io import BytesIO

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from app.core.files_config import BRAND_CONFIG
from app.domain.entities.document_content import DocumentContent
from app.domain.interfaces.document_builder import IDocumentBuilder
from app.services.document_generation.template_engine import TemplateEngine

logger = logging.getLogger(__name__)


def set_cell_background(cell, fill_hex):
    """Establece el color de fondo de una celda de tabla."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def clear_table_borders(table):
    """Elimina todos los bordes de una tabla."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_borders(table, color_hex="DBDFFC", sz="4"):
    """Aplica bordes finos a una tabla."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_margins(table, top=100, bottom=100, left=150, right=150):
    """Establece márgenes internos (padding) para las celdas de una tabla."""
    tblPr = table._tbl.tblPr
    cellMar = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tblPr.append(cellMar)


def add_xml_field_to_run(run, field_name):
    """Inserta un campo dinámico de Word (p. ej. PAGE, NUMPAGES) en un run."""
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> {field_name} </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


def format_run(run, font_name="Arial", size_pt=10.5, color_rgb=None, bold=False, italic=False):
    """Aplica formato básico a un objeto run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb


def format_paragraph(paragraph, space_before_pt=0, space_after_pt=6, line_spacing=1.15, keep_with_next=False):
    """Aplica formato de espaciado y flujo a un párrafo."""
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(space_before_pt)
    p_format.space_after = Pt(space_after_pt)
    p_format.line_spacing = line_spacing
    p_format.keep_with_next = keep_with_next


def number_sections(sections):
    """Genera títulos numerados jerárquicamente para las secciones."""
    h1_count = 0
    h2_count = 0
    h3_count = 0
    numbered_list = []
    
    for s in sections:
        if s.level == 1:
            h1_count += 1
            h2_count = 0
            h3_count = 0
            num_str = f"{h1_count}. "
        elif s.level == 2:
            h2_count += 1
            h3_count = 0
            num_str = f"{h1_count}.{h2_count}. "
        else:
            h3_count += 1
            num_str = f"{h1_count}.{h2_count}.{h3_count}. "
            
        title = s.title.strip()
        # Evitar duplicar si ya empieza con un número
        if not re.match(r'^\d+(\.\d+)*\.', title):
            numbered_title = num_str + title
        else:
            numbered_title = title
        numbered_list.append((s, numbered_title, num_str))
        
    return numbered_list


def add_grey_badge(doc, text):
    """Agrega una pequeña etiqueta/badge gris antes del título."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    
    # Ancho fijo del badge de 1.4 pulgadas
    table.columns[0].width = Inches(1.4)
    cell = table.cell(0, 0)
    cell.width = Inches(1.4)
    
    set_cell_background(cell, "F5F5F5")
    clear_table_borders(table)
    set_table_margins(table, top=40, bottom=40, left=100, right=100)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before_pt=0, space_after_pt=0)
    run = p.add_run(text)
    format_run(run, font_name="Arial", size_pt=8, color_rgb=RGBColor(113, 128, 150), bold=True)


class DocxBuilder(IDocumentBuilder):

    def __init__(self, engine: TemplateEngine):
        self._engine = engine

    @property
    def output_format(self) -> str:
        return "docx"

    def build(self, content: DocumentContent) -> bytes:
        brand = content.brand or "convertia"
        cfg = BRAND_CONFIG.get(brand, BRAND_CONFIG["convertia"])

        try:
            # Crear documento desde cero
            doc = docx.Document()

            # Configuración de página (Tamaño Carta y márgenes de 1 pulgada)
            section = doc.sections[0]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

            # Diferenciar la primera página para omitir encabezado y pie de página en la portada
            section.different_first_page_header_footer = True

            # -------------------------------------------------------------
            # Configurar Encabezado (Páginas Interiores: Solo Isotopo)
            # -------------------------------------------------------------
            logo_docs_path = cfg["logos"].get("docs")  # Isotopo circular
            if logo_docs_path and os.path.exists(logo_docs_path):
                header = section.header
                hp = header.paragraphs[0]
                hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                hrun = hp.add_run()
                hrun.add_picture(logo_docs_path, height=Inches(0.35))

            # -------------------------------------------------------------
            # Configurar Pie de Página (Páginas Interiores: Copyright y Pág.)
            # -------------------------------------------------------------
            footer = section.footer
            footer_table = footer.add_table(1, 2, Inches(6.5))
            footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            clear_table_borders(footer_table)

            cell_left = footer_table.cell(0, 0)
            cell_left.width = Inches(4.5)
            p_left = cell_left.paragraphs[0]
            format_paragraph(p_left, space_before_pt=0, space_after_pt=0)
            r_left = p_left.add_run("© Intelligence Customer Acquisition — Convertia — Documentación interna")
            format_run(r_left, font_name="Arial", size_pt=8, color_rgb=RGBColor(113, 128, 150))

            cell_right = footer_table.cell(0, 1)
            cell_right.width = Inches(2.0)
            p_right = cell_right.paragraphs[0]
            format_paragraph(p_right, space_before_pt=0, space_after_pt=0)
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            r_pfx = p_right.add_run()
            format_run(r_pfx, font_name="Arial", size_pt=8, color_rgb=RGBColor(113, 128, 150))
            add_xml_field_to_run(r_pfx, "PAGE")

            # -------------------------------------------------------------
            # PORTADA MINIMALISTA (SEGÚN MOCKUP)
            # -------------------------------------------------------------
            # Logotipo principal arriba a la derecha (convertia_main.png)
            p_cover_logo = doc.add_paragraph()
            format_paragraph(p_cover_logo, space_before_pt=0, space_after_pt=120)
            p_cover_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo_main_path = cfg["logos"].get("main")
            if logo_main_path and os.path.exists(logo_main_path):
                r_logo = p_cover_logo.add_run()
                r_logo.add_picture(logo_main_path, height=Inches(0.35))

           

            # Título principal en dos tonos (última palabra en gris)
            p_title = doc.add_paragraph()
            format_paragraph(p_title, space_before_pt=18, space_after_pt=8)
            
            words = content.title.split()
            if len(words) > 1:
                part1 = " ".join(words[:-1]) + " "
                part2 = words[-1]
            else:
                part1 = content.title
                part2 = ""

            r_title1 = p_title.add_run(part1)
            format_run(r_title1, font_name="Arial", size_pt=32, color_rgb=RGBColor(1, 30, 35), bold=True)

            if part2:
                r_title2 = p_title.add_run(part2)
                format_run(r_title2, font_name="Arial", size_pt=32, color_rgb=RGBColor(113, 128, 150), bold=True)

            # Subtítulo (color verde corporativo oscuro)
            p_sub = doc.add_paragraph()
            format_paragraph(p_sub, space_before_pt=4, space_after_pt=0)
            if content.subtitle:
                r_sub = p_sub.add_run(content.subtitle)
                format_run(r_sub, font_name="Arial", size_pt=13, color_rgb=RGBColor(16, 71, 63))

            # Salto de página para el Índice
            doc.add_page_break()

            # Numerar y mapear páginas para el Índice
            numbered_sections = number_sections(content.sections)
            
            page_map = {}
            current_page = 1  # La portada es 1 en Word físicamente, pero el índice se cuenta y muestra 1
            # Para estimar correctamente:
            # - Página 1: Portada (oculta número)
            # - Página 1: Índice (muestra 1)
            # - Página 2: Comienzo de secciones
            section_page = 2
            for idx, (s_item, num_title, num_str) in enumerate(numbered_sections):
                if s_item.level == 1 and idx > 0:
                    section_page += 1
                page_map[id(s_item)] = section_page

            # Mapeo de páginas de las tablas globales (se colocan al final en una nueva página)
            total_sections_pages = section_page
            table_page = total_sections_pages + 1

            # -------------------------------------------------------------
            # DIAPOSITIVA DE ÍNDICE (TOC)
            # -------------------------------------------------------------
            p_idx_title = doc.add_paragraph()
            format_paragraph(p_idx_title, space_before_pt=24, space_after_pt=18)
            run_idx_title = p_idx_title.add_run("ÍNDICE")
            format_run(run_idx_title, font_name="Arial", size_pt=16, color_rgb=RGBColor(1, 30, 35), bold=True)

            for s_item, num_title, num_str in numbered_sections:
                p_item = doc.add_paragraph()
                format_paragraph(p_item, space_before_pt=4, space_after_pt=4)
                
                # Configurar tabulación con puntos líderes en el margen derecho (6.5 in)
                p_item.paragraph_format.tab_stops.add_tab_stop(
                    Inches(6.5),
                    alignment=docx.enum.text.WD_TAB_ALIGNMENT.RIGHT,
                    leader=docx.enum.text.WD_TAB_LEADER.DOTS
                )

                # Sangría según nivel
                if s_item.level == 2:
                    p_item.paragraph_format.left_indent = Inches(0.25)
                elif s_item.level == 3:
                    p_item.paragraph_format.left_indent = Inches(0.5)

                p_num = page_map[id(s_item)]
                r_item = p_item.add_run(f"{num_title}\t{p_num}")
                format_run(r_item, font_name="Arial", size_pt=10, color_rgb=RGBColor(45, 55, 72))

            # Salto de página tras el índice
            doc.add_page_break()

            # -------------------------------------------------------------
            # DIAPOSITIVAS DE CONTENIDO (SECCIONES Y TABLAS)
            # -------------------------------------------------------------
            for s_idx, (section_item, numbered_title, num_str) in enumerate(numbered_sections):
                level = section_item.level

                if level == 1:
                    if s_idx > 0:
                        doc.add_page_break()

                    p = doc.add_paragraph()
                    format_paragraph(p, space_before_pt=24, space_after_pt=18, keep_with_next=True)
                    r = p.add_run(numbered_title)
                    format_run(r, font_name="Arial", size_pt=16, color_rgb=RGBColor(1, 30, 35), bold=True)

                elif level == 2:
                    p = doc.add_paragraph()
                    format_paragraph(p, space_before_pt=18, space_after_pt=8, keep_with_next=True)
                    r = p.add_run(numbered_title)
                    format_run(r, font_name="Arial", size_pt=13, color_rgb=RGBColor(1, 30, 35), bold=True)

                else:
                    p = doc.add_paragraph()
                    format_paragraph(p, space_before_pt=14, space_after_pt=6, keep_with_next=True)
                    r = p.add_run(numbered_title)
                    format_run(r, font_name="Arial", size_pt=11.5, color_rgb=RGBColor(45, 55, 72), bold=True)

                # Párrafos de texto
                if section_item.content:
                    for para in section_item.content.split("\n"):
                        if para.strip():
                            if para.strip().startswith(">"):
                                # Renderizar como bloque de cita/destacado
                                text = para.strip().lstrip(">").strip()
                                self._add_docx_callout_box(doc, text)
                            else:
                                p_text = doc.add_paragraph()
                                format_paragraph(p_text, space_before_pt=0, space_after_pt=6)
                                p_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                r_text = p_text.add_run(para.strip())
                                format_run(r_text, font_name="Arial", size_pt=10.5, color_rgb=RGBColor(45, 55, 72))

                # Viñetas
                if section_item.bullets:
                    for bullet in section_item.bullets:
                        if bullet.strip():
                            p_bullet = doc.add_paragraph(style="List Bullet")
                            format_paragraph(p_bullet, space_before_pt=0, space_after_pt=4)
                            r_bullet = p_bullet.add_run(bullet.strip())
                            format_run(r_bullet, font_name="Arial", size_pt=10.5, color_rgb=RGBColor(45, 55, 72))

                # Tabla inline
                if section_item.table:
                    self._add_docx_table(doc, section_item.table)

            # Tablas globales al final del documento
            if content.tables:
                doc.add_page_break()
                p_table_title = doc.add_paragraph()
                format_paragraph(p_table_title, space_before_pt=24, space_after_pt=18)
                # Título de la sección de tablas
                r_tab_title = p_table_title.add_run(f"{len(content.sections)+1}. Ejemplo de tabla")
                format_run(r_tab_title, font_name="Arial", size_pt=16, color_rgb=RGBColor(1, 30, 35), bold=True)

                for table in content.tables:
                    self._add_docx_table(doc, table)

            # Guardar documento en buffer
            buffer = BytesIO()
            doc.save(buffer)
            logger.info(f"DOCX generado con diseño de mockup: '{content.title}'")
            return buffer.getvalue()

        except Exception as e:
            logger.error(f"Error generando DOCX mockup '{content.title}': {e}")
            raise RuntimeError(f"Error al generar el archivo DOCX: {e}") from e

    def _add_docx_callout_box(self, doc, text):
        """Dibuja una caja de llamado premium oscura (#011E23) con texto en negrita y acentos verdes."""
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(6.5)

        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "011E23")
        clear_table_borders(table)
        set_table_margins(table, top=160, bottom=160, left=200, right=200)

        p = cell.paragraphs[0]
        format_paragraph(p, space_before_pt=4, space_after_pt=4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parsear marcas de negrita '**' para aplicar verde menta
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                format_run(run, font_name="Arial", size_pt=10.5, color_rgb=RGBColor(26, 235, 159), bold=True)
            else:
                run = p.add_run(part)
                format_run(run, font_name="Arial", size_pt=10.5, color_rgb=RGBColor(255, 255, 255), bold=True)

        # Espacio tras la caja
        p_space = doc.add_paragraph()
        format_paragraph(p_space, space_before_pt=0, space_after_pt=10)

    def _add_docx_table(self, doc, table_data):
        headers = table_data.headers
        rows = table_data.rows
        caption = table_data.caption

        if not headers:
            return

        if caption:
            p_cap = doc.add_paragraph()
            format_paragraph(p_cap, space_before_pt=12, space_after_pt=4, keep_with_next=True)
            r_cap = p_cap.add_run(caption.upper())
            format_run(r_cap, font_name="Arial", size_pt=9, color_rgb=RGBColor(113, 128, 150), bold=True)

        num_rows = len(rows) + 1
        num_cols = len(headers)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        set_table_borders(table, color_hex="DBDFFC", sz="4")
        set_table_margins(table, top=100, bottom=100, left=150, right=150)

        col_width = Inches(6.5 / num_cols)
        for col in table.columns:
            col.width = col_width

        # Escribir Cabecera (#011E23)
        hdr_cells = table.rows[0].cells
        for c_idx, header in enumerate(headers):
            cell = hdr_cells[c_idx]
            cell.width = col_width
            set_cell_background(cell, "011E23")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            format_paragraph(p, space_before_pt=2, space_after_pt=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(header))
            format_run(r, font_name="Arial", size_pt=9.5, color_rgb=RGBColor(255, 255, 255), bold=True)

        # Escribir Datos (Cebra o Destacado)
        for r_idx, row in enumerate(rows):
            row_cells = table.rows[r_idx + 1].cells
            
            # Detectar fila destacada (si contiene la palabra "Destacado" en cualquier celda)
            is_highlight = any(str(val).strip().lower() == "destacado" for val in row)
            
            if is_highlight:
                bg_hex = "E6FFFA"  # Fondo cian claro
                text_color = RGBColor(16, 71, 63)  # Texto verde bosque oscuro
                is_bold = True
            else:
                bg_hex = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
                text_color = RGBColor(45, 55, 72)
                is_bold = False

            for c_idx, val in enumerate(row):
                if c_idx >= num_cols:
                    break
                cell = row_cells[c_idx]
                cell.width = col_width
                set_cell_background(cell, bg_hex)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                format_paragraph(p, space_before_pt=2, space_after_pt=2)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(str(val))
                format_run(r, font_name="Arial", size_pt=9.5, color_rgb=text_color, bold=is_bold)

        p_space = doc.add_paragraph()
        format_paragraph(p_space, space_before_pt=0, space_after_pt=10)
