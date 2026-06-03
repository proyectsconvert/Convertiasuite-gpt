
import logging
from io import BytesIO
from docx import Document as DocxDocument
from app.domain.entities.document import DocumentType, ParsedContent, Section, Table
from app.domain.interfaces.document_processor import IDocumentProcessor

logger = logging.getLogger(__name__)


class DocxProcessor(IDocumentProcessor):

    @property
    def supported_type(self) -> DocumentType:
        return DocumentType.DOCX

    @property
    def supported_extensions(self) -> list[str]:
        return ["docx", "doc"]

    def get_metadata(self, file_content: bytes) -> dict:
        try:
            doc = DocxDocument(BytesIO(file_content))
            props = doc.core_properties
            return {
                "author": props.author or "",
                "title": props.title or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
            }
        except Exception as e:
            logger.warning(f"No se pudieron extraer metadatos del DOCX: {e}")
            return {}

    async def parse(self, file_content: bytes, filename: str) -> ParsedContent:
        try:
            text = self.extract_text(file_content)
            tables = self._extract_tables(file_content)
            sections = self._extract_sections(file_content)

            return ParsedContent(
                text=text,
                sections=sections,
                tables=tables,
                metadata=self.get_metadata(file_content),
            )
        except Exception as e:
            logger.error(f"Error parseando DOCX '{filename}': {e}")
            raise ValueError(f"Error al procesar el archivo Word (.docx): {e}")

    def _extract_tables(self, file_content: bytes) -> list[Table]:
        doc = DocxDocument(BytesIO(file_content))
        tables = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append(
                    [cell.text.strip().replace("\n", " ") for cell in row.cells]
                )
            if rows:
                headers = rows[0]
                data_rows = [
                    [str(cell) for cell in row] for row in rows[1:]
                ]
                tables.append(
                    Table(
                        name=f"Tabla {i + 1}",
                        headers=headers,
                        rows=data_rows,
                    )
                )
        return tables

    def _extract_sections(self, file_content: bytes) -> list[Section]:
        doc = DocxDocument(BytesIO(file_content))
        sections = []
        current_title = "Introducción"
        current_level = 1
        current_content: list[str] = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            if style_name.startswith("Heading"):
                if current_content:
                    sections.append(
                        Section(
                            title=current_title,
                            level=current_level,
                            content="\n".join(current_content),
                        )
                    )
                    current_content = []

                try:
                    current_level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    current_level = 1
                current_title = text
            else:
                current_content.append(text)

        if current_content:
            sections.append(
                Section(
                    title=current_title,
                    level=current_level,
                    content="\n".join(current_content),
                )
            )

        return sections

    @staticmethod
    def extract_text(file_content: bytes) -> str:
        try:
            doc = DocxDocument(BytesIO(file_content))
            elements = []

            for element in doc.element.body:
                if element.tag.endswith("p"):
                    for p in doc.paragraphs:
                        if p._element is element:
                            if p.text and p.text.strip():
                                elements.append(p.text.strip())
                            break
                elif element.tag.endswith("tbl"):
                    for table in doc.tables:
                        if table._element is element:
                            table_md = []
                            for row_idx, row in enumerate(table.rows):
                                row_cells = [
                                    cell.text.strip().replace("\n", " ")
                                    for cell in row.cells
                                ]
                                table_md.append("| " + " | ".join(row_cells) + " |")
                                if row_idx == 0:
                                    table_md.append(
                                        "| " + " | ".join(["---"] * len(row_cells)) + " |"
                                    )
                            if table_md:
                                elements.append("\n".join(table_md))
                            break

            if not elements:
                return "El archivo Word (.docx) está vacío."

            return "\n\n".join(elements)

        except Exception as e:
            logger.error(f"Error parseando DOCX: {e}")
            raise ValueError(f"Error al procesar el archivo Word (.docx): {e}")
