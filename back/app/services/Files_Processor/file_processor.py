from io import BytesIO
import logging
import pandas as pd
import json
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class FileProcessorService:
    @staticmethod
    def extract_text_from_pdf(contents: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(contents))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(f"--- Página {i + 1} ---\n{text.strip()}")

            if not pages_text:
                return "El PDF parece estar vacío o no se pudo extraer texto. Verifique si contiene solo imágenes."

            return "\n\n".join(pages_text)
        except Exception as e:
            logger.error(f"Error parseando PDF: {str(e)}")
            raise ValueError(f"Error al procesar el archivo PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(contents: bytes) -> str:
        try:
            doc = Document(BytesIO(contents))
            elements = []

            for element in doc.element.body:
                if element.tag.endswith("p"):
                    for p in doc.paragraphs:
                        if p._element is element:
                            if p.text and p.text.strip():
                                elements.append(p.text.strip())
                            break
                elif element.tag.endswith("tbl"):
                    for table in doc.tables:
                        if table._element is element:
                            table_md = []
                            for row_idx, row in enumerate(table.rows):
                                row_cells = [
                                    cell.text.strip().replace("\n", " ")
                                    for cell in row.cells
                                ]
                                table_md.append("| " + " | ".join(row_cells) + " |")
                                if row_idx == 0:
                                    table_md.append(
                                        "| "
                                        + " | ".join(["---"] * len(row_cells))
                                        + " |"
                                    )
                            if table_md:
                                elements.append("\n".join(table_md))
                            break

            if not elements:
                for p in doc.paragraphs:
                    if p.text and p.text.strip():
                        elements.append(p.text.strip())
                for table in doc.tables:
                    table_md = []
                    for row_idx, row in enumerate(table.rows):
                        row_cells = [
                            cell.text.strip().replace("\n", " ") for cell in row.cells
                        ]
                        table_md.append("| " + " | ".join(row_cells) + " |")
                        if row_idx == 0:
                            table_md.append(
                                "| " + " | ".join(["---"] * len(row_cells)) + " |"
                            )
                    if table_md:
                        elements.append("\n".join(table_md))

            if not elements:
                return "El archivo Word (.docx) está vacío."

            return "\n\n".join(elements)
        except Exception as e:
            logger.error(f"Error parseando archivo DOCX: {str(e)}")
            raise ValueError(f"Error al procesar el archivo Word (.docx): {str(e)}")

    @staticmethod
    def _analizar_reporte_convertia(
        df: pd.DataFrame, cols_normalizadas: dict
    ) -> str | None:
        col_mkt = next(
            (
                cols_normalizadas[c]
                for c in cols_normalizadas
                if "CAMPAÑA_MKT" in c or "CAMPANA_MKT" in c
            ),
            None,
        )
        col_llave = next(
            (
                cols_normalizadas[c]
                for c in cols_normalizadas
                if "ID_LLAVE" in c or "IDLLAVE" in c
            ),
            None,
        )

        if not col_mkt or not col_llave:
            return None

        try:
            total_leads = len(df)

            def check_is_sale(val):
                if pd.isna(val):
                    return False
                s_val = str(val).strip().lower()
                return s_val not in ["", "nan", "null", "none"]

            df["__ES_VENTA"] = df[col_llave].apply(check_is_sale)
            total_ventas = int(df["__ES_VENTA"].sum())

            leads_por_campana = df[col_mkt].value_counts().head(10).to_dict()
            ventas_por_campana = (
                df[df["__ES_VENTA"] == True][col_mkt].value_counts().head(10).to_dict()
            )

            conversion_por_campana = {}
            for campana in leads_por_campana:
                total_c_leads = len(df[df[col_mkt] == campana])
                total_c_ventas = len(
                    df[(df[col_mkt] == campana) & (df["__ES_VENTA"] == True)]
                )
                rate = (
                    round((total_c_ventas / total_c_leads) * 100, 2)
                    if total_c_leads > 0
                    else 0
                )
                conversion_por_campana[campana] = (
                    f"{rate}% ({total_c_ventas} ventas de {total_c_leads} leads)"
                )

            bi_payload = {
                "REPORTE_EJECUTIVO_BI_CONVERTIA": {
                    "metodologia_aplicada": "Filtrado estricto donde la columna ID_LLAVE posee un identificador de venta válido.",
                    "resumen_global": {
                        "total_leads_recibidos": total_leads,
                        "total_ventas_exitosas": total_ventas,
                        "tasa_conversion_global": f"{round((total_ventas / total_leads) * 100, 2) if total_leads > 0 else 0}%",
                    },
                    "top_campanas_por_volumen_leads": leads_por_campana,
                    "top_campanas_por_ventas_efectivas": ventas_por_campana,
                    "eficiencia_y_conversion_por_campana": conversion_por_campana,
                }
            }
            return (
                f"--- PROCESAMIENTO AUTOMÁTICO DE DATOS COMERCIALES ---\n"
                f"Usa este resumen estructurado con las métricas reales del archivo completo para responder las preguntas de BI:\n"
                f"{json.dumps(bi_payload, indent=2, ensure_ascii=False)}"
            )
        except Exception as e:
            logger.error(f"Error en análisis especializado BI Convertia: {str(e)}")
            return None

    @staticmethod
    def _analizar_reporte_financiero(
        df: pd.DataFrame, cols_normalizadas: dict
    ) -> str | None:
        col_ingresos = next(
            (
                cols_normalizadas[c]
                for c in cols_normalizadas
                if "INGRESOS" in c or "REVENUE" in c
            ),
            None,
        )
        if not col_ingresos:
            return None
        return "--- ANALISIS DE EXPERTO FINANCIERO (Pendiente implementar) ---"

    @classmethod
    def analyze_tabular_data(cls, rows: list[list]) -> str:
        if not rows or len(rows) < 2:
            return "El archivo está vacío o no contiene filas válidas."

        try:
            df_check = pd.DataFrame(rows[1:], columns=rows[0])
            cols_normalizadas = {
                str(col).strip().upper(): col for col in df_check.columns
            }

            analisis_mkt = cls._analizar_reporte_convertia(df_check, cols_normalizadas)
            if analisis_mkt:
                return analisis_mkt

            analisis_fin = cls._analizar_reporte_financiero(df_check, cols_normalizadas)
            if analisis_fin:
                return analisis_fin

            total_filas = len(df_check)
            columnas_reales = [
                str(c).strip() for c in df_check.columns if c is not None
            ]

            if total_filas <= 150:
                datos_completos = df_check.to_dict(orient="records")

                payload_generico = {
                    "METADATA_DEL_ARCHIVO": {
                        "tipo_dataset": "Documento Estructurado General (Completo)",
                        "total_registros": total_filas,
                        "columnas_disponibles": columnas_reales,
                    },
                    "CONTENIDO_REAL_FILA_POR_FILA": datos_completos,
                }
                return (
                    f"--- ARCHIVO PROCESADO EXITOSAMENTE ---\n"
                    f"Analiza la siguiente base de datos completa para responder con precisión matemática la solicitud del usuario:\n"
                    f"{json.dumps(payload_generico, indent=2, ensure_ascii=False)}"
                )

            else:
                muestra_filas = df_check.head(25).to_dict(orient="records")

                payload_grande = {
                    "METADATA_DEL_ARCHIVO": {
                        "tipo_dataset": "Dataset de Gran Volumen (Muestra Acotada)",
                        "total_registros": total_filas,
                        "columnas_disponibles": columnas_reales,
                        "valores_unicos_por_columna": {
                            str(col): int(df_check[col].nunique())
                            for col in df_check.columns
                        },
                    },
                    "VISTA_PREVIA_REGISTROS_REALES": muestra_filas,
                }
                return (
                    f"--- DATASET GRANDE PROCESADO ---\n"
                    f"Usa los metadatos y esta muestra representativa para responder al usuario. "
                    f"Si te pide un cruce masivo que requiera más datos, descríbele los campos disponibles en base a este payload:\n"
                    f"{json.dumps(payload_grande, indent=2, ensure_ascii=False)}"
                )

        except Exception as e:
            logger.warning(
                f"No se pudo ejecutar el pre-analisis optimizado, cayendo en flujo clásico: {e}"
            )

        raw_headers = rows[0]
        headers = [
            str(h).strip()
            if h is not None and str(h).strip() != ""
            else f"Columna_{i + 1}"
            for i, h in enumerate(raw_headers)
        ]
        data_rows = rows[1:]
        total_rows = len(data_rows)
        num_cols = len(headers)

        summary = [
            "### Resumen General del Dataset (Fallback)",
            f"- **Total de filas de datos**: {total_rows}",
            f"- **Total de columnas**: {num_cols}",
            f"- **Nombres de columnas**: {', '.join([f'`{h}`' for h in headers])}",
            "",
        ]

        summary.append("### Vista Previa de Filas Completas")
        summary.append("| " + " | ".join(headers) + " |")
        summary.append("| " + " | ".join(["---"] * num_cols) + " |")
        for r in data_rows[:10]:
            row_cells = [
                str(r[col_idx]).strip()
                if col_idx < len(r) and r[col_idx] is not None
                else ""
                for col_idx in range(num_cols)
            ]
            summary.append("| " + " | ".join(row_cells) + " |")

        return "\n".join(summary)

    @classmethod
    def extract_text_from_csv(cls, contents: bytes) -> str:
        try:
            try:
                df = pd.read_csv(
                    BytesIO(contents), sep=None, engine="python", encoding="utf-8"
                )
            except Exception:
                df = pd.read_csv(
                    BytesIO(contents), sep=None, engine="python", encoding="latin-1"
                )

            df = df.fillna("")
            headers = list(df.columns)
            data_rows = df.values.tolist()
            rows = [headers] + data_rows

            return cls.analyze_tabular_data(rows)
        except Exception as e:
            logger.error(f"Error parseando CSV con Pandas: {str(e)}")
            raise ValueError(f"Error al procesar el archivo CSV: {str(e)}")

    @classmethod
    def extract_text_from_excel(cls, contents: bytes) -> str:
        try:
            df = pd.read_excel(BytesIO(contents))

            df = df.fillna("")
            headers = list(df.columns)
            data_rows = df.values.tolist()
            rows = [headers] + data_rows

            return cls.analyze_tabular_data(rows)
        except Exception as e:
            logger.error(f"Error parseando Excel con Pandas: {str(e)}")
            raise ValueError(f"Error al procesar el archivo Excel: {str(e)}")

    @staticmethod
    def extract_text_from_txt(contents: bytes) -> str:
        try:
            try:
                text = contents.decode("utf-8")
            except UnicodeDecodeError:
                text = contents.decode("latin-1")
                if not text or not text.strip():
                    return "El archivo de texto parece estar vacío."
            return text.strip()
        except Exception as e:
            logger.error(f"Error parseando archivo de texto: {str(e)}")
            raise ValueError(f"Error al procesar el archivo de texto: {str(e)}")

    @staticmethod
    def extract_text_from_json(contents: bytes) -> str:
        try:
            text_contend = contents.decode("utf-8")
            if not text_contend.strip():
                return "El archivo JSON parece estar vacío."
            parsed_json = json.loads(text_contend)
            return json.dumps(parsed_json, indent=2, ensure_ascii=False)
        except UnicodeDecodeError:
            logger.error("Error decodificando JSON")
            raise ValueError(
                "Error al decodificar el archivo JSON. Verifique la codificación del archivo."
            )
        except json.JSONDecodeError as e:
            logger.error(f"Error parseando JSON: {str(e)}")
            raise ValueError(f"Error al procesar el archivo JSON: {str(e)}")

    @staticmethod
    def extract_text_from_md(contents: bytes) -> str:
        try:
            try:
                text = contents.decode("utf-8")
            except UnicodeDecodeError:
                text = contents.decode("latin-1")

            if not text or not text.strip():
                return "El archivo Markdown (.md) está vacío."

            return text.strip()
        except Exception as e:
            logger.error(f"Error parseando archivo MD: {str(e)}")
            raise ValueError(f"Error al procesar el archivo Markdown (.md): {str(e)}")

    @staticmethod
    def extract_base64_from_image(contents: bytes) -> str:
        import base64
        try:
            return base64.b64encode(contents).decode("utf-8")
        except Exception as e:
            logger.error(f"Error codificando imagen a Base64: {str(e)}")
            raise ValueError(f"Error al procesar la imagen: {str(e)}")


class FileProcessorFactory:
    @staticmethod
    def get_parser(filename: str, content_type: str = None) -> tuple:
        filename_lower = filename.lower()
        if filename_lower.endswith((".png", ".jpg", ".jpeg", ".webp")) or (content_type and content_type.startswith("image/")):
            return FileProcessorService.extract_base64_from_image, "image"
        elif filename_lower.endswith(".pdf") or content_type == "application/pdf":
            return FileProcessorService.extract_text_from_pdf, "pdf"
        elif (
            filename_lower.endswith((".docx", ".doc"))
            or content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return FileProcessorService.extract_text_from_docx, "word"
        elif filename_lower.endswith(".csv") or content_type == "text/csv":
            return FileProcessorService.extract_text_from_csv, "csv"
        elif (
            filename_lower.endswith((".xlsx", ".xls"))
            or content_type
            == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ):
            return FileProcessorService.extract_text_from_excel, "excel"
        elif filename_lower.endswith((".txt", ".text")) or content_type == "text/plain":
            return FileProcessorService.extract_text_from_txt, "txt"
        elif filename_lower.endswith(".json") or content_type == "application/json":
            return FileProcessorService.extract_text_from_json, "json"
        elif (
            filename_lower.endswith((".md", ".markdown"))
            or content_type == "text/markdown"
        ):
            return FileProcessorService.extract_text_from_md, "md"
        return None, None
