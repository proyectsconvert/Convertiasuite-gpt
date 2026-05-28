import io
import json
import logging
import re
import pandas as pd
from typing import Any
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily
from reportlab.platypus import (
    PageBreak,
    Spacer,
    SimpleDocTemplate,
    Paragraph,
    Table as PDFTable,
    TableStyle,
)
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import urllib.request
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class FileGeneratorService:
    @staticmethod
    def generate_txt(content: str) -> bytes:
        return content.encode("utf-8")

    @staticmethod
    def generate_md(content: str) -> bytes:
        return content.encode("utf-8")

    @staticmethod
    def generate_json(content: Any) -> bytes:
        try:
            if isinstance(content, str):
                parsed = json.loads(content)
                return json.dumps(parsed, indent=2, ensure_ascii=False).encode("utf-8")
            return json.dumps(content, indent=2, ensure_ascii=False).encode("utf-8")
        except Exception as e:
            logger.error(f"Error generating JSON: {e}")
            if isinstance(content, str):
                return content.encode("utf-8")
            return str(content).encode("utf-8")

    @classmethod
    def generate_csv(cls, content: Any) -> bytes:
        try:
            df = cls._convert_to_dataframe(content)
            output = io.BytesIO()
            df.to_csv(output, index=False, encoding="utf-8")
            return output.getvalue()
        except Exception as e:
            logger.error(f"Error generating CSV: {e}")
            if isinstance(content, str):
                return content.encode("utf-8")
            raise ValueError(f"Failed to generate CSV: {str(e)}")

    @classmethod
    def generate_excel(cls, content: Any) -> bytes:
        try:
            df = cls._convert_to_dataframe(content)
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                df.to_excel(writer, index=False, sheet_name="Datos")
            return output.getvalue()
        except Exception as e:
            logger.error(f"Error generating Excel: {e}")
            raise ValueError(f"Failed to generate Excel: {str(e)}")

    @staticmethod
    def generate_docx(content: str) -> bytes:
        if Document is None:
            raise ImportError("python-docx is not installed.")

        try:
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(0.88)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            font_name = "Inter"
            c_verde_turquesa = "0CD78E"
            c_azul_oscuro = RGBColor(1, 30, 35)
            c_verde_oscuro = RGBColor(16, 71, 63)
            c_gris_oscuro = RGBColor(13, 13, 14)
            c_petroleo_oscuro = "002422"
            c_blanco_lavanda = "F4F6F6"
            c_gris_lavanda = "DBDFFC"
            c_text_body = RGBColor(45, 55, 72)
            c_meta = RGBColor(74, 85, 104)
            c_table_text = RGBColor(26, 32, 44)

            def add_formatted_para(
                text_content,
                font_size,
                rgb_color,
                space_before=0,
                space_after=0,
                align=0,
                bold=False,
                style=None,
            ):
                if style:
                    p = doc.add_paragraph(style=style)
                else:
                    p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                p.paragraph_format.line_spacing = 1.2
                if align == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                parts = re.split(r"(<b>.*?</b>|<i>.*?</i>)", text_content)
                for part in parts:
                    if not part:
                        continue
                    run = p.add_run()
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = rgb_color
                    if part.startswith("<b>") and part.endswith("</b>"):
                        run.text = part[3:-4]
                        run.bold = True
                    elif part.startswith("<i>") and part.endswith("</i>"):
                        run.text = part[3:-4]
                        run.italic = True
                    else:
                        run.text = part
                        run.bold = bold
                return p

            lines = content.split("\n")
            detected_title = "Reporte Corporativo"
            for line in lines:
                clean_raw = line.strip().replace("**", "").replace("#", "").strip()
                if line.strip().startswith("# ") and clean_raw:
                    detected_title = clean_raw
                    break
                elif "Reporte de" in line or "Análisis" in line:
                    detected_title = line.replace('"', "").replace("**", "").strip()
                    break

            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(140)

            add_formatted_para(
                FileGeneratorService.clean_markdown_tags(detected_title),
                40,
                c_azul_oscuro,
                space_after=15,
                bold=True,
            )
            add_formatted_para(
                "ConvertGPT— Reporte generado por Olivia AI",
                22,
                c_gris_oscuro,
                space_after=15,
            )

            p_space2 = doc.add_paragraph()
            p_space2.paragraph_format.space_before = Pt(180)

            fecha_actual = datetime.now().strftime("%d / %m / %Y")
            add_formatted_para(
                f"<b>Fecha de generación:</b> {fecha_actual}", 10, c_meta, space_after=6
            )
            add_formatted_para(
                "<b>Clasificación:</b> Confidencial Corporativo",
                10,
                c_meta,
                space_after=6,
            )

            doc.add_page_break()

            section_body = doc.sections[0]
            footer = section_body.footer
            footer_p = footer.paragraphs[0]
            footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            footer_p.paragraph_format.space_before = Pt(10)
            footer_p._p.get_or_add_pPr().append(
                parse_xml(
                    f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="{c_verde_turquesa}"/></w:pBdr>'
                )
            )

            run_f = footer_p.add_run(
                "©  Intelligence Customer Acquisition — Convertia — Documentación interna"
            )
            run_f.font.name = font_name
            run_f.font.size = Pt(8)
            run_f.font.color.rgb = RGBColor(73, 77, 82)

            run_f2 = footer_p.add_run("\t\tPágina ")
            run_f2.font.name = font_name
            run_f2.font.size = Pt(8)
            run_f2.font.color.rgb = RGBColor(73, 77, 82)
            footer_p._p.append(
                parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>')
            )

            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                if stripped.replace("#", "").strip() == detected_title and i < 5:
                    i += 1
                    continue

                if stripped.startswith("|"):
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith("|"):
                        table_lines.append(lines[i].strip())
                        i += 1

                    table_data = []
                    for t_line in table_lines:
                        parts = [p.strip() for p in t_line.split("|")]
                        if parts and parts[0] == "":
                            parts = parts[1:]
                        if parts and parts[-1] == "":
                            parts = parts[:-1]
                        if all(
                            all(c == "-" or c == ":" or c == " " for c in cell)
                            for cell in parts
                            if cell
                        ):
                            continue
                        if parts:
                            table_data.append(parts)

                    if table_data:
                        num_rows = len(table_data)
                        num_cols = len(table_data[0])
                        w_table = doc.add_table(rows=num_rows, cols=num_cols)
                        w_table.autofit = True

                        for r_idx, row_cells in enumerate(table_data):
                            for c_idx, cell_value in enumerate(row_cells):
                                w_cell = w_table.cell(r_idx, c_idx)
                                cell_cleaned = FileGeneratorService.clean_markdown_tags(
                                    cell_value
                                )
                                p_cell = w_cell.paragraphs[0]
                                p_cell.paragraph_format.space_before = Pt(4)
                                p_cell.paragraph_format.space_after = Pt(4)
                                p_cell.paragraph_format.line_spacing = 1.1

                                parts = re.split(
                                    r"(<b>.*?</b>|<i>.*?</i>)", cell_cleaned
                                )
                                for part in parts:
                                    if not part:
                                        continue
                                    run = p_cell.add_run()
                                    run.font.name = font_name
                                    run.font.size = Pt(8.5)
                                    if r_idx == 0:
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                        if part.startswith("<b>") and part.endswith(
                                            "</b>"
                                        ):
                                            run.text = part[3:-4]
                                            run.bold = True
                                        elif part.startswith("<i>") and part.endswith(
                                            "</i>"
                                        ):
                                            run.text = part[3:-4]
                                            run.italic = True
                                        else:
                                            run.text = part
                                            run.bold = True
                                    else:
                                        run.font.color.rgb = c_table_text
                                        if part.startswith("<b>") and part.endswith(
                                            "</b>"
                                        ):
                                            run.text = part[3:-4]
                                            run.bold = True
                                        elif part.startswith("<i>") and part.endswith(
                                            "</i>"
                                        ):
                                            run.text = part[3:-4]
                                            run.italic = True
                                        else:
                                            run.text = part

                                if r_idx == 0:
                                    shd = parse_xml(
                                        f'<w:shd {nsdecls("w")} w:fill="{c_petroleo_oscuro}"/>'
                                    )
                                Gold_bg = (
                                    c_blanco_lavanda if r_idx % 2 == 1 else "FFFFFF"
                                )
                                if r_idx > 0:
                                    shd = parse_xml(
                                        f'<w:shd {nsdecls("w")} w:fill="{Gold_bg}"/>'
                                    )
                                w_cell._tc.get_or_add_tcPr().append(shd)
                                w_cell._tc.get_or_add_tcPr().append(
                                    parse_xml(
                                        f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="{c_gris_lavanda}"/><w:left w:val="single" w:sz="4" w:space="0" w:color="{c_gris_lavanda}"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="{c_gris_lavanda}"/><w:right w:val="single" w:sz="4" w:space="0" w:color="{c_gris_lavanda}"/></w:tcBorders>'
                                    )
                                )

                        doc.add_paragraph().paragraph_format.space_after = Pt(12)
                    continue

                cleaned_line = FileGeneratorService.clean_markdown_tags(stripped)

                if stripped.startswith("# "):
                    add_formatted_para(
                        cleaned_line[2:], 22, c_gris_oscuro, space_after=15, bold=True
                    )
                elif stripped.startswith("## "):
                    add_formatted_para(
                        cleaned_line[3:],
                        18,
                        c_verde_oscuro,
                        space_before=16,
                        space_after=8,
                        bold=True,
                    )
                elif stripped.startswith("### "):
                    add_formatted_para(
                        cleaned_line[4:],
                        14,
                        RGBColor(45, 55, 72),
                        space_before=12,
                        space_after=6,
                        bold=True,
                    )
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    add_formatted_para(
                        cleaned_line[2:],
                        11,
                        c_text_body,
                        space_after=5,
                        style="List Bullet",
                    )
                elif stripped.startswith("1. ") or (
                    stripped
                    and stripped[0].isdigit()
                    and "." in stripped.split(" ", 1)[0]
                ):
                    dot_idx = cleaned_line.find(".")
                    text_part = cleaned_line[dot_idx + 1 :].strip()
                    add_formatted_para(
                        text_part, 9.5, c_text_body, space_after=5, style="List Number"
                    )
                elif stripped:
                    add_formatted_para(cleaned_line, 9.5, c_text_body, space_after=8)
                else:
                    p_empty = doc.add_paragraph()
                    p_empty.paragraph_format.space_after = Pt(4)
                i += 1

            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise ValueError(f"Failed to generate DOCX: {str(e)}")

    @staticmethod
    def clean_markdown_tags(text: str) -> str:
        text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
        text = re.sub(r"__(.*?)__", r"<b>\1</b>", text)
        text = re.sub(r"_(.*?)_", r"<i>\1</i>", text)
        if text.startswith('"') and text.endswith('"') and text.count('"') == 2:
            text = text[1:-1]
        return text

    @classmethod
    def generate_pdf(cls, content: str) -> bytes:
        if letter is None:
            raise ImportError("reportlab is not installed.")

        try:
            font_normal = "inter"
            font_bold = "inter-bold"

            if "inter" not in pdfmetrics.getRegisteredFontNames():
                try:
                    font_dir = (
                        os.path.dirname(__file__) if "__file__" in locals() else "."
                    )
                    reg_path = os.path.join(font_dir, "Inter-Regular.ttf")
                    bold_path = os.path.join(font_dir, "Inter-Bold.ttf")
                    if not os.path.exists(reg_path):
                        urllib.request.urlretrieve(
                            "https://cdn.jsdelivr.net/gh/rsms/inter@v4.0/docs/font-files/Inter-Regular.ttf",
                            reg_path,
                        )
                    if not os.path.exists(bold_path):
                        urllib.request.urlretrieve(
                            "https://cdn.jsdelivr.net/gh/rsms/inter@v4.0/docs/font-files/Inter-Bold.ttf",
                            bold_path,
                        )
                    pdfmetrics.registerFont(TTFont("inter", reg_path))
                    pdfmetrics.registerFont(TTFont("inter-bold", bold_path))
                    registerFontFamily("inter", normal="inter", bold="inter-bold")
                except Exception as font_err:
                    logger.warning(
                        f"No se pudo cargar Inter, usando Helvetica por defecto: {font_err}"
                    )
                    font_normal = "Helvetica"
                    font_bold = "Helvetica-Bold"

            pdf_buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=letter,
                rightMargin=54,
                leftMargin=54,
                topMargin=72,
                bottomMargin=64,
            )

            styles = getSampleStyleSheet()
            color_verde_turquesa = colors.HexColor("#0cd78e")
            color_azul_oscuro = colors.HexColor("#011e23")
            color_verde_oscuro = colors.HexColor("#10473F")
            color_gris_oscuro = colors.HexColor("#0D0D0E")
            color_petroleo_oscuro = colors.HexColor("#002422")
            color_blanco_lavanda = colors.HexColor("#F4F6F6")
            color_gris_lavanda = colors.HexColor("#DBDFFC")

            COVER_TITLE_STYLE = ParagraphStyle(
                "CoverTitle",
                parent=styles["Title"],
                fontName=font_bold,
                fontSize=40,
                leading=54,
                textColor=color_azul_oscuro,
                spaceAfter=15,
                alignment=0,
            )
            cover_subtitle_style = ParagraphStyle(
                "CoverSubtitle",
                parent=styles["Normal"],
                fontName=font_normal,
                fontSize=22,
                leading=22,
                textColor=color_gris_oscuro,
                spaceAfter=15,
                alignment=0,
            )
            cover_meta_style = ParagraphStyle(
                "CoverMeta",
                parent=styles["Normal"],
                fontName=font_normal,
                fontSize=10,
                leading=16,
                textColor=colors.HexColor("#4A5568"),
                spaceAfter=6,
            )
            title_style = ParagraphStyle(
                "PDFTitle",
                parent=styles["Title"],
                fontName=font_bold,
                fontSize=22,
                leading=26,
                textColor=color_gris_oscuro,
                spaceAfter=15,
                alignment=0,
            )
            h1_style = ParagraphStyle(
                "PDFH1",
                parent=styles["Heading1"],
                fontName=font_bold,
                fontSize=18,
                leading=22,
                textColor=color_verde_oscuro,
                spaceBefore=16,
                spaceAfter=8,
                keepWithNext=True,
            )
            h2_style = ParagraphStyle(
                "PDFH2",
                parent=styles["Heading2"],
                fontName=font_bold,
                fontSize=14,
                leading=20,
                textColor=colors.HexColor("#2D3748"),
                spaceBefore=12,
                spaceAfter=6,
                keepWithNext=True,
            )
            body_style = ParagraphStyle(
                "PDFBody",
                parent=styles["BodyText"],
                fontName=font_normal,
                fontSize=11,
                leading=15,
                textColor=colors.HexColor("#2D3748"),
                spaceAfter=8,
            )
            bullet_style = ParagraphStyle(
                "PDFBullet",
                parent=body_style,
                leftIndent=15,
                firstLineIndent=-10,
                spaceAfter=5,
            )

            story = []
            lines = content.split("\n")

            for line in lines:
                clean_raw = line.strip().replace("**", "").replace("#", "").strip()
                if line.strip().startswith("# ") and clean_raw:
                    detected_title = clean_raw
                    break
                elif "Reporte de" in line or "Análisis" in line:
                    detected_title = line.replace('"', "").replace("**", "").strip()
                    break

            story.append(Spacer(1, 140))
            story.append(
                Paragraph(cls.clean_markdown_tags(detected_title), COVER_TITLE_STYLE)
            )
            story.append(
                Paragraph(
                    "ConvertGPT— Reporte generado por Olivia AI", cover_subtitle_style
                )
            )

            line_decorativa = PDFTable([[""]], colWidths=[120], rowHeights=[3.5])
            line_decorativa.setStyle(
                TableStyle([("BACKGROUND", (0, 0), (-1, -1), color_verde_turquesa)])
            )
            story.append(line_decorativa)

            story.append(Spacer(1, 180))

            fecha_actual = datetime.now().strftime("%d / %m / %Y")
            story.append(
                Paragraph(
                    f"<b>Fecha de generación:</b> {fecha_actual}", cover_meta_style
                )
            )
            story.append(
                Paragraph(
                    "<b>Clasificación:</b> Confidencial Corporativo", cover_meta_style
                )
            )

            story.append(PageBreak())

            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                if stripped.replace("#", "").strip() == detected_title and i < 5:
                    i += 1
                    continue

                if stripped.startswith("|"):
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith("|"):
                        table_lines.append(lines[i].strip())
                        i += 1

                    table_data = []
                    for t_line in table_lines:
                        parts = [p.strip() for p in t_line.split("|")]
                        if parts and parts[0] == "":
                            parts = parts[1:]
                        if parts and parts[-1] == "":
                            parts = parts[:-1]
                        if all(
                            all(c == "-" or c == ":" or c == " " for c in cell)
                            for cell in parts
                            if cell
                        ):
                            continue
                        if parts:
                            table_data.append(parts)

                    if table_data:
                        formatted_table_data = []
                        for row_idx, row in enumerate(table_data):
                            formatted_row = []
                            for cell in row:
                                cell_cleaned = cls.clean_markdown_tags(cell)
                                cell_style = ParagraphStyle(
                                    f"Cell_{row_idx}",
                                    parent=body_style,
                                    fontSize=8.5,
                                    leading=11,
                                    textColor=colors.HexColor("#1A202C")
                                    if row_idx > 0
                                    else colors.white,
                                    fontName=font_bold if row_idx == 0 else font_normal,
                                )
                                formatted_row.append(
                                    Paragraph(cell_cleaned, cell_style)
                                )
                            formatted_table_data.append(formatted_row)

                        col_count = len(table_data[0])
                        col_widths = (
                            [doc.width / col_count] * col_count if col_count else None
                        )
                        pdf_table = PDFTable(formatted_table_data, colWidths=col_widths)
                        pdf_table.setStyle(
                            TableStyle(
                                [
                                    (
                                        "BACKGROUND",
                                        (0, 0),
                                        (-1, 0),
                                        color_petroleo_oscuro,
                                    ),
                                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                                    ("GRID", (0, 0), (-1, -1), 0.5, color_gris_lavanda),
                                    (
                                        "ROWBACKGROUNDS",
                                        (0, 1),
                                        (-1, -1),
                                        [colors.white, color_blanco_lavanda],
                                    ),
                                ]
                            )
                        )
                        story.append(pdf_table)
                        story.append(Spacer(1, 12))
                    continue

                cleaned_line = cls.clean_markdown_tags(stripped)

                if stripped.startswith("# "):
                    story.append(Paragraph(cleaned_line[2:], title_style))
                elif stripped.startswith("## "):
                    story.append(Paragraph(cleaned_line[3:], h1_style))
                elif stripped.startswith("### "):
                    story.append(Paragraph(cleaned_line[4:], h2_style))
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    story.append(Paragraph(f"• {cleaned_line[2:]}", bullet_style))
                elif stripped.startswith("1. ") or (
                    stripped
                    and stripped[0].isdigit()
                    and "." in stripped.split(" ", 1)[0]
                ):
                    story.append(Paragraph(cleaned_line, bullet_style))
                elif stripped:
                    story.append(Paragraph(cleaned_line, body_style))
                else:
                    story.append(Spacer(1, 4))
                i += 1

            def draw_cover_background(canvas, document):
                canvas.saveState()
                canvas.restoreState()

            def draw_header_footer(canvas, document):
                canvas.saveState()
                canvas.setStrokeColor(color_verde_turquesa)
                canvas.setLineWidth(0.5)
                canvas.line(54, 45, letter[0] - 54, 45)
                canvas.setFillColor(colors.HexColor("#494D52"))
                canvas.setFont(
                    font_normal
                    if font_normal in pdfmetrics.getRegisteredFontNames()
                    else "Helvetica",
                    8,
                )
                canvas.drawString(
                    54,
                    30,
                    "©  Intelligence Customer Acquisition — Convertia — Documentación interna",
                )
                canvas.drawRightString(letter[0] - 54, 30, f"Página {document.page}")
                canvas.restoreState()

            doc.build(
                story,
                onFirstPage=draw_cover_background,
                onLaterPages=draw_header_footer,
            )
            return pdf_buffer.getvalue()
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            raise ValueError(f"Failed to generate PDF: {str(e)}")

    @staticmethod
    def _convert_to_dataframe(content: Any) -> pd.DataFrame:
        if isinstance(content, pd.DataFrame):
            return content
        if isinstance(content, str):
            try:
                parsed = json.loads(content)
                if isinstance(parsed, list):
                    return pd.DataFrame(parsed)
                return pd.DataFrame([parsed])
            except Exception:
                try:
                    return pd.read_csv(io.StringIO(content))
                except Exception:
                    return pd.DataFrame([{"Contenido": content}])
        if isinstance(content, list):
            if all(isinstance(row, list) for row in content) and len(content) > 0:
                headers = content[0]
                rows = content[1:]
                return pd.DataFrame(rows, columns=headers)
            return pd.DataFrame(content)
        if isinstance(content, dict):
            return pd.DataFrame([content])
        return pd.DataFrame([{"Datos": str(content)}])
